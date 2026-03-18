from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)

# ── Helper: set paragraph style ───────────────────────────────
def set_para_fmt(para, space_before=0, space_after=8, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)

def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        if level == 1:
            run.font.size  = Pt(20)
            run.font.color.rgb = RGBColor(0, 86, 210)
            run.font.bold  = True
        elif level == 2:
            run.font.size  = Pt(14)
            run.font.color.rgb = RGBColor(28, 29, 31)
            run.font.bold  = True
        elif level == 3:
            run.font.size  = Pt(12)
            run.font.color.rgb = RGBColor(80, 80, 80)
            run.font.bold  = True
    set_para_fmt(para, space_before=18, space_after=6)
    return para

def add_body(doc, text, italic=False):
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.size   = Pt(11)
        run.font.italic = italic
        run.font.color.rgb = RGBColor(60, 60, 60)
    set_para_fmt(para, space_before=2, space_after=6, line_spacing=14)
    return para

def add_bullet(doc, text, indent=0):
    para = doc.add_paragraph(style='List Bullet')
    run  = para.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(60, 60, 60)
    para.paragraph_format.left_indent = Inches(0.3 + indent * 0.3)
    set_para_fmt(para, space_before=1, space_after=3)
    return para

def add_hr(doc):
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_fmt(para, space_before=4, space_after=4)
    return para

def add_kv(doc, key, value):
    para = doc.add_paragraph()
    r1 = para.add_run(f"{key}: ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(28, 29, 31)
    r2 = para.add_run(value)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(80, 80, 80)
    set_para_fmt(para, space_before=2, space_after=4)
    return para

# ══════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_para.add_run("CoursIDLLM")
r.font.size  = Pt(36)
r.font.bold  = True
r.font.color.rgb = RGBColor(0, 86, 210)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Product & Engineering Document")
r2.font.size  = Pt(16)
r2.font.color.rgb = RGBColor(100, 100, 100)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run(f"Version 1.0  ·  {datetime.date.today().strftime('%B %d, %Y')}")
r3.font.size  = Pt(11)
r3.font.color.rgb = RGBColor(160, 160, 160)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  1. PRODUCT REQUIREMENTS DOCUMENT
# ══════════════════════════════════════════════════════════════
add_heading(doc, "1. Product Requirements Document", 1)
add_hr(doc)

add_heading(doc, "1.1 Overview", 2)
add_body(doc,
    "CoursIDLLM is an AI-augmented online learning platform that combines structured "
    "course delivery with LLM-powered tutoring, skills gap analysis, and career path "
    "recommendations. The platform targets individuals, businesses, universities, and "
    "government agencies seeking scalable, personalised education.")

add_heading(doc, "1.2 Goals", 2)
for g in [
    "Provide a world-class catalogue of video-based courses across tech, design, and business domains.",
    "Integrate an LLM assistant that answers learner questions in the context of the active course.",
    "Deliver a Skills Builder tool that maps current skills to learning recommendations.",
    "Deliver a Career Builder tool that aligns course completions to job-market opportunities.",
    "Support multi-language content and UI (English, Arabic, French, Spanish, Chinese).",
    "Serve four distinct audience segments with tailored landing experiences.",
]:
    add_bullet(doc, g)

add_heading(doc, "1.3 Non-Goals (v1)", 2)
for ng in [
    "Live instructor-led sessions (planned for v2).",
    "Mobile native apps (responsive web first).",
    "Custom LMS white-labelling for enterprise clients.",
]:
    add_bullet(doc, ng)

add_heading(doc, "1.4 Key Features", 2)
features = {
    "Two-tier Navigation": "Audience bar (Individuals / Businesses / Universities / Governments) + main nav with Courses dropdown, Skills Builder, Career Builder, search, language switcher, Log In, Join for Free.",
    "Course Catalogue": "Filterable by category, level, price (free/paid), rating, and language. Rich course detail pages with preview video, syllabus, and instructor bio.",
    "LLM Tutor": "Per-lesson AI chat powered by Claude API. Context-aware answers scoped to course content.",
    "Skills Builder": "Interactive skill-gap assessment → personalised learning path.",
    "Career Builder": "Job role matching based on completed courses and declared goals.",
    "Authentication": "Email/password + OAuth (Google, GitHub). JWT sessions.",
    "Payments": "Stripe integration for individual course purchase and subscription plans.",
    "Dashboard": "Learner progress tracking, certificates, wishlist, and LLM chat history.",
    "Admin Panel": "Course management, user analytics, revenue reports.",
}
for title, desc in features.items():
    para = doc.add_paragraph()
    r = para.add_run(f"{title}: ")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0, 86, 210)
    r2 = para.add_run(desc)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(60, 60, 60)
    set_para_fmt(para, space_before=3, space_after=5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  2. TECH STACK
# ══════════════════════════════════════════════════════════════
add_heading(doc, "2. Tech Stack", 1)
add_hr(doc)

layers = {
    "Frontend": [
        ("Framework",   "Next.js 15 (App Router, React Server Components)"),
        ("Styling",     "Tailwind CSS v4"),
        ("Language",    "TypeScript"),
        ("State",       "Zustand (client) + React Query (server state)"),
        ("Forms",       "React Hook Form + Zod validation"),
        ("Video",       "Video.js or Mux Player"),
        ("i18n",        "next-intl"),
    ],
    "Backend / API": [
        ("Runtime",     "Node.js 22 via Next.js API Routes / Route Handlers"),
        ("ORM",         "Prisma"),
        ("Auth",        "NextAuth.js v5 (JWT + OAuth providers)"),
        ("Payments",    "Stripe SDK"),
        ("Storage",     "AWS S3 / Cloudflare R2 for video & assets"),
        ("Email",       "Resend (transactional emails)"),
    ],
    "AI / LLM": [
        ("Model",       "Anthropic Claude (claude-sonnet-4-6 default)"),
        ("SDK",         "Anthropic TypeScript SDK (@anthropic-ai/sdk)"),
        ("Streaming",   "Vercel AI SDK for streaming responses to client"),
        ("RAG",         "Pinecone vector DB for course-content embeddings"),
    ],
    "Database": [
        ("Primary DB",  "PostgreSQL (Supabase hosted)"),
        ("Cache",       "Redis (Upstash) for sessions and rate limiting"),
        ("Search",      "Meilisearch for full-text course search"),
    ],
    "DevOps / Infra": [
        ("Hosting",     "Vercel (Next.js) + Railway (background workers)"),
        ("CI/CD",       "GitHub Actions"),
        ("Monitoring",  "Sentry (errors) + Vercel Analytics"),
        ("Secrets",     ".env.local + Vercel environment variables"),
    ],
}

for layer, items in layers.items():
    add_heading(doc, f"2.{list(layers.keys()).index(layer)+1} {layer}", 2)
    for key, val in items:
        add_kv(doc, key, val)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  3. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════
add_heading(doc, "3. Project Structure", 1)
add_hr(doc)

add_body(doc, "Next.js App Router layout with src/ directory:", italic=True)

structure = """\
coursidllm/
├── src/
│   ├── app/
│   │   ├── (marketing)/          # Public pages (/, /about, /pricing)
│   │   │   └── page.tsx
│   │   ├── (auth)/               # Login, signup, forgot-password
│   │   │   ├── login/page.tsx
│   │   │   └── signup/page.tsx
│   │   ├── catalog/              # Course catalogue + search
│   │   │   └── page.tsx
│   │   ├── course/[slug]/        # Course detail page
│   │   │   └── page.tsx
│   │   ├── learn/[slug]/         # In-course player + LLM chat
│   │   │   ├── page.tsx
│   │   │   └── ChatPanel.tsx
│   │   ├── dashboard/            # Learner dashboard
│   │   │   └── page.tsx
│   │   ├── skills/               # Skills Builder
│   │   │   └── page.tsx
│   │   ├── careers/              # Career Builder
│   │   │   └── page.tsx
│   │   ├── admin/                # Admin panel (protected)
│   │   │   └── page.tsx
│   │   ├── api/
│   │   │   ├── auth/[...nextauth]/route.ts
│   │   │   ├── courses/route.ts
│   │   │   ├── chat/route.ts     # LLM streaming endpoint
│   │   │   ├── payments/route.ts
│   │   │   └── webhooks/stripe/route.ts
│   │   ├── layout.tsx            # Root layout (SiteHeader + SiteFooter)
│   │   └── globals.css
│   ├── components/
│   │   ├── SiteHeader.tsx        # Two-tier sticky header ✅
│   │   ├── SiteFooter.tsx
│   │   ├── CourseCard.tsx
│   │   ├── CourseGrid.tsx
│   │   ├── SearchBar.tsx
│   │   ├── VideoPlayer.tsx
│   │   ├── LLMChat.tsx
│   │   ├── SkillsAssessment.tsx
│   │   └── ui/                  # Reusable primitives (Button, Badge, Modal…)
│   ├── lib/
│   │   ├── prisma.ts             # Prisma client singleton
│   │   ├── anthropic.ts          # Anthropic SDK client
│   │   ├── auth.ts               # NextAuth config
│   │   ├── stripe.ts             # Stripe client
│   │   └── utils.ts
│   ├── hooks/
│   │   ├── useCourseProgress.ts
│   │   └── useLLMStream.ts
│   ├── store/                    # Zustand slices
│   │   ├── cartStore.ts
│   │   └── userStore.ts
│   └── types/
│       └── index.ts
├── prisma/
│   └── schema.prisma
├── public/
├── .env.local
├── next.config.ts
├── tailwind.config.ts
└── package.json"""

code_para = doc.add_paragraph()
run = code_para.add_run(structure)
run.font.name = "Courier New"
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(30, 30, 30)
code_para.paragraph_format.space_before = Pt(4)
code_para.paragraph_format.space_after  = Pt(4)
code_para.paragraph_format.left_indent  = Inches(0.3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  4. SCHEMA DESIGN
# ══════════════════════════════════════════════════════════════
add_heading(doc, "4. Schema Design", 1)
add_hr(doc)
add_body(doc, "PostgreSQL via Prisma ORM. Key models:", italic=True)

schemas = {
    "User": ["id (uuid)", "name", "email (unique)", "passwordHash", "avatarUrl",
             "role (STUDENT | INSTRUCTOR | ADMIN)", "language", "createdAt", "updatedAt"],
    "Course": ["id (uuid)", "slug (unique)", "title", "description", "thumbnailUrl",
               "previewVideoUrl", "price (Decimal)", "isFree (bool)", "level (BEGINNER | INTERMEDIATE | ADVANCED)",
               "category", "language", "published (bool)", "instructorId → User",
               "createdAt", "updatedAt"],
    "Lesson": ["id", "courseId → Course", "title", "videoUrl", "duration (seconds)",
               "order (int)", "content (Markdown transcript)"],
    "Enrollment": ["id", "userId → User", "courseId → Course", "enrolledAt",
                   "completedAt (nullable)", "progressPercent"],
    "LessonProgress": ["id", "enrollmentId → Enrollment", "lessonId → Lesson",
                       "watchedSeconds", "completed (bool)", "updatedAt"],
    "Order": ["id", "userId → User", "stripePaymentIntentId", "amount (Decimal)",
              "currency", "status (PENDING | PAID | REFUNDED)", "createdAt"],
    "OrderItem": ["id", "orderId → Order", "courseId → Course", "price (Decimal)"],
    "Review": ["id", "userId → User", "courseId → Course", "rating (1-5)", "comment", "createdAt"],
    "ChatMessage": ["id", "userId → User", "lessonId → Lesson", "role (USER | ASSISTANT)",
                    "content (Text)", "createdAt"],
    "SkillAssessment": ["id", "userId → User", "skills (JSON)", "gaps (JSON)",
                        "recommendedCourses (JSON)", "createdAt"],
    "Certificate": ["id", "userId → User", "courseId → Course", "issuedAt", "certificateUrl"],
}

for model, fields in schemas.items():
    add_heading(doc, model, 3)
    for f in fields:
        add_bullet(doc, f)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  5. USER FLOW
# ══════════════════════════════════════════════════════════════
add_heading(doc, "5. User Flow", 1)
add_hr(doc)

flows = {
    "5.1 Guest → Enrolment": [
        "Land on homepage (/) → see Hero + featured courses.",
        "Click audience tab (e.g. For Businesses) → hero content updates.",
        "Use search bar or Courses dropdown → arrive at /catalog.",
        "Click course card → /course/[slug] (detail page).",
        'Click "Enrol Now" → redirect to /login (if not authenticated).',
        "Sign up (/signup) or log in (/login) → redirect back to course detail.",
        'Complete Stripe checkout → enrolled; redirect to /learn/[slug].',
    ],
    "5.2 Learning Flow": [
        "Dashboard (/dashboard) → lists enrolled courses and progress.",
        "Click 'Continue' → /learn/[slug] opens video player.",
        "Lesson completes → LessonProgress updated, next lesson auto-plays.",
        "LLM Chat panel available per lesson for Q&A on course content.",
        "All lessons complete → certificate generated; shown on /dashboard.",
    ],
    "5.3 Skills Builder Flow": [
        "Navigate to /skills.",
        "Complete interactive skills self-assessment quiz.",
        "Platform calculates skill gaps vs. desired role.",
        "Receive personalised course recommendations → links to /catalog.",
    ],
    "5.4 Career Builder Flow": [
        "Navigate to /careers.",
        "Declare target job role or industry.",
        "Platform maps completed courses + skills to job requirements.",
        "Show missing skills and recommended courses to close the gap.",
        "Show live job postings (future integration via external Job API).",
    ],
    "5.5 Admin Flow": [
        "Admin logs in → /admin dashboard.",
        "Manage courses: create, edit, publish/unpublish.",
        "View learner analytics: enrollments, completion rates, revenue.",
        "Manage users: roles, bans, certificate revocation.",
    ],
}

for title, steps in flows.items():
    add_heading(doc, title, 2)
    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(style='List Number')
        run  = para.add_run(step)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(60, 60, 60)
        set_para_fmt(para, space_before=1, space_after=3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  6. STYLING GUIDELINES
# ══════════════════════════════════════════════════════════════
add_heading(doc, "6. Styling Guidelines", 1)
add_hr(doc)

add_heading(doc, "6.1 Design Principles", 2)
for p in [
    "Clean & professional — no heavy gradients on main UI, reserved for hero/CTA sections.",
    "Accessible — WCAG 2.1 AA minimum contrast ratios on all text.",
    "Consistent spacing — 4 px base unit; spacing scale: 4, 8, 12, 16, 24, 32, 48, 64 px.",
    "Mobile-first — all layouts designed for 375 px, scaled up.",
]:
    add_bullet(doc, p)

add_heading(doc, "6.2 Colour Palette", 2)
colours = [
    ("Primary Blue",   "#0056d2", "Main CTAs, logo, links, active states"),
    ("Purple Accent",  "#a435f0", "Badges, highlights, hover accents"),
    ("Dark Surface",   "#1c1d1f", "Audience bar, hero, footer background"),
    ("Body Text",      "#3d3d3d", "Main paragraph text"),
    ("Muted Text",     "#6a6f73", "Secondary labels, captions"),
    ("Border",         "#d1d7dc", "Input borders, dividers, card borders"),
    ("Background",     "#f7f9fa", "Alternating section backgrounds"),
    ("White",          "#ffffff", "Card surfaces, nav background"),
    ("Success",        "#1e7e34", "Completion states, positive indicators"),
    ("Error",          "#c0392b", "Validation errors, destructive actions"),
]
for name, hex_val, usage in colours:
    para = doc.add_paragraph()
    r1 = para.add_run(f"{name} ({hex_val}): ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(28, 29, 31)
    r2 = para.add_run(usage)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(80, 80, 80)
    set_para_fmt(para, space_before=2, space_after=4)

add_heading(doc, "6.3 Typography", 2)
type_rules = [
    ("Font family",      "Inter (system-ui fallback in offline environments)"),
    ("Heading 1",        "32–40 px / font-weight 800 / line-height 1.15"),
    ("Heading 2",        "24–28 px / font-weight 700"),
    ("Heading 3",        "18–20 px / font-weight 600"),
    ("Body",             "15–16 px / font-weight 400 / line-height 1.65"),
    ("Small / caption",  "12–13 px / font-weight 500"),
    ("Code / mono",      "Courier New or JetBrains Mono, 13 px"),
]
for k, v in type_rules:
    add_kv(doc, k, v)

add_heading(doc, "6.4 Component Tokens (Tailwind)", 2)
tokens = [
    ("Border radius — buttons", "rounded (4 px)"),
    ("Border radius — cards",   "rounded-xl (12 px)"),
    ("Border radius — pills",   "rounded-full (99 px)"),
    ("Shadow — card",           "shadow-md"),
    ("Shadow — dropdown",       "shadow-xl"),
    ("Transition",              "duration-150 ease-in-out on all interactive elements"),
    ("Focus ring",              "ring-2 ring-blue-200 on all inputs and buttons"),
]
for k, v in tokens:
    add_kv(doc, k, v)

add_heading(doc, "6.5 Header-Specific Rules", 2)
for rule in [
    "Audience bar height: 36 px / bg #1c1d1f / active link: white text + 2 px white bottom border.",
    "Main nav height: 56 px / bg white / sticky top-0 z-50 / bottom border #d1d7dc.",
    "Logo: #0056d2 / 22 px / font-weight 800. Accent letter in #a435f0.",
    "Search bar: border #1c1d1f default → #0056d2 on focus. Search button bg #0056d2.",
    "Language switcher: outlined button (border #d1d7dc) with globe icon + language code.",
    "Log In: plain text link. Join for Free: outlined button with #0056d2 border, fills on hover.",
    "Dropdowns: rounded-xl, shadow-xl, appear on click, close on outside click.",
]:
    add_bullet(doc, rule)

# ── Save ──────────────────────────────────────────────────────
output = "/home/user/CoursIDLLM_ProductDoc.docx"
doc.save(output)
print(f"Saved: {output}")
